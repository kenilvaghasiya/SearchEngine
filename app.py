from flask import Flask, request, jsonify
import os
from pathlib import Path
from JsonFileDocument import JsonFileDocument
import PyPDF2
from XmlHtmlDocument import XmlHtmlDocument
from TokenProcessor import TokenProcessor,spanishToken,frenchToken
from indexing import PositionalInvertedIndex
from docx import Document 
from langdetect import detect
from querying import BooleanQueryParser
import sys
from textblob import TextBlob
from flask_cors import CORS


app = Flask(__name__)
CORS(app, origins=["http://localhost:3000","http://localhost:3000"])
documents = {}
directories = ["Data/json", "Data/txt", "Data/pdf", "Data/xml", "Data/docx"]
jsonFile=PositionalInvertedIndex()
txtFile=PositionalInvertedIndex()
pdfFile=PositionalInvertedIndex()
xmlFile=PositionalInvertedIndex()
docxFile=PositionalInvertedIndex()




boolean_query_parser=None

def load_pdf(pdf_file_path):
    with open(pdf_file_path, 'rb') as pdf_file:
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        num_pages = len(pdf_reader.pages)  # Use len(reader.pages) to get the number of pages
        body = ""

        for page_num in range(num_pages):
            page = pdf_reader.pages[page_num]
            body += page.extract_text()

        title = os.path.splitext(os.path.basename(pdf_file_path))[0] # Use the file name as the title
        return {"title":title, "body":body, 'filetype':'pdf'}

def LoadDocuments(directory_path):
    documents = []
    for filename in os.listdir(directory_path):
        file_path = os.path.join(directory_path, filename)
        if filename.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as txt_file:
                title = os.path.splitext(filename)[0]  # Use the file name as the title
                content = txt_file.read()
                documents.append({'title': title, 'body': content,'filetype':'txt'})
        elif filename.endswith('.json'):
            json_document = JsonFileDocument(file_path)
            documents.append({'title': file_path.split('/')[-1]+" "+json_document.get_title(), 'body': json_document.get_body(), 'url': json_document.get_url(),'filetype':'json'})
        elif filename.endswith(('.xml', '.html', '.htm')):
            xml_html_document = XmlHtmlDocument(file_path)
            # print(xml_html_document.filetype)
            documents.append(xml_html_document.dataSend())
        elif filename.endswith('.pdf'):
            pdf_document = load_pdf(file_path)
            documents.append(pdf_document)
        elif filename.endswith('.docx'):
            # Handle .docx files using python-docx
            doc = Document(file_path)
            text_content = " ".join([paragraph.text for paragraph in doc.paragraphs])
            documents.append({'title': os.path.splitext(os.path.basename(file_path))[0], 'body': text_content, 'filetype':'docx'})
    return documents

def GetTokenData(rowdata):
    # Check if rowdata['body'] is empty or None
    if not rowdata['body']:
        return {"fileName": rowdata["title"], "tokenData": []}

    detected_language = detect(rowdata['body'])
    print(f"\n\nThe detected language is: {detected_language}\n\n") 

    if detected_language == "en":
        token_processor = TokenProcessor()
        resulting_types = token_processor.process_token(rowdata['body'])
        resulting_terms = [token_processor.normalize_type(t) for t in resulting_types]
        # filtered_terms = token_processor.remove_stopwords(resulting_types)
        return {"fileName": rowdata["title"], "tokenData": resulting_terms , "indexAndToken":token_processor.get_tokens_and_positions(rowdata['body'])}
    
    elif detected_language == "es":
        processor = spanishToken()
        spanish_tokens, spanish_terms = processor.process_text(rowdata['body'])
        return {"fileName": rowdata["title"], "tokenData": spanish_tokens,"indexAndToken":[]}
    elif detected_language == "fr":
        processor = frenchToken()
        fr_tokens, fr_terms = processor.process_text(rowdata['body'])
        return {"fileName": rowdata["title"], "tokenData": fr_tokens,"indexAndToken":[]}
    else:
        return {"fileName": rowdata["title"], "tokenData": [],"indexAndToken":[]}



def load_files():
    for i in range(len(directories)):
        documents[directories[i]]=LoadDocuments(directories[i])
    for document_id, document in enumerate(documents["Data/txt"]):
        data=GetTokenData(document)
        print(len(data["tokenData"]),len(data['indexAndToken']))
        for token in data["tokenData"]:
            txtFile.add_term(token, data["fileName"])
        for tokens in data['indexAndToken']:
            txtFile.add_termIndex(tokens,data["fileName"])
    for document_id, document in enumerate(documents["Data/pdf"]):
        data=GetTokenData(document)
        print(len(data["tokenData"]),len(data['indexAndToken']))
        for token in data["tokenData"]:
            pdfFile.add_term(token, data["fileName"])
        for tokens in data['indexAndToken']:
            pdfFile.add_termIndex(tokens,data["fileName"])
    for document_id, document in enumerate(documents["Data/xml"]):
        data=GetTokenData(document)
        print(len(data["tokenData"]),len(data['indexAndToken']))
        for token in data["tokenData"]:
            xmlFile.add_term(token, data["fileName"])
        for tokens in data['indexAndToken']:
            xmlFile.add_termIndex(tokens,data["fileName"])
    for document_id, document in enumerate(documents["Data/docx"]):
        data=GetTokenData(document)
        print(len(data["tokenData"]),len(data['indexAndToken']))
        for token in data["tokenData"]:
            docxFile.add_term(token, data["fileName"])
        for tokens in data['indexAndToken']:
            docxFile.add_termIndex(tokens,data["fileName"])
    for document_id, document in enumerate(documents["Data/json"]):
        data=GetTokenData(document)
        print(len(data["tokenData"]),len(data['indexAndToken']))
        for token in data["tokenData"]:
            jsonFile.add_term(token, data["fileName"])
        for tokens in data['indexAndToken']:
            jsonFile.add_termIndex(tokens,data["fileName"])
    
    return jsonify({'message': documents["Data/txt"]})



with app.app_context():
        load_files() 
   


    
@app.route('/searchdata', methods=['POST'])
def process_string():
    try:
        # Get the JSON data from the request
        boolean_query_parser = BooleanQueryParser(jsonFile)
        data = request.get_json()
        print(data)
        if 'text' not in data:
            return jsonify({'error': 'Missing "text" field in JSON data'}), 400
        if 'type' not in data:
            return jsonify({'error': 'Missing "text" field in JSON data'}), 400
       
        if data['type']=="json":
          boolean_query_parser = BooleanQueryParser(jsonFile)
        if data['type']=="txt":
            boolean_query_parser = BooleanQueryParser(txtFile)
        if data['type']=="pdf":
            boolean_query_parser = BooleanQueryParser(pdfFile)
        if data['type']=="xml":
            boolean_query_parser = BooleanQueryParser(xmlFile)   
        if data['type']=="docx":
            boolean_query_parser = BooleanQueryParser(docxFile)
            
        # Split the received string (assuming it's space-separated)
        text = data['text']
        parsed_query = boolean_query_parser.parse_query(text)
        postings = boolean_query_parser.get_postings(parsed_query)

        if not postings:
            response_data = {'data':[],"message":"Term not found in any documents","file":0}
        else:
            response_data = {'data':postings,"message":"done","file":len(postings)}
        return jsonify(response_data), 200

    except Exception as e:
        return jsonify({'error': str(e)}), 500



